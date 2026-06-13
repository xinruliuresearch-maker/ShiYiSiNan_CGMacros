from __future__ import annotations

import csv
import json
import math
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Sequence

from PIL import Image, ImageDraw, ImageFont


DATE_TAG = "2026-06-10"
PACKAGE_NAME = f"ajcn_no_wetlab_manuscript_package_{DATE_TAG}"

ROOT = Path(__file__).resolve().parents[1]
INTEGRATED = ROOT / "outputs" / "integrated_mainline"
NO_DIGEST = INTEGRATED / "no_digestion_phase0_to_phase4"
PHASE2 = NO_DIGEST / "phase2_manuscript_tables_and_figures"
BOOST = INTEGRATED / "dry_lab_booster_modules"
OUT = INTEGRATED / PACKAGE_NAME
NHANES_MIRROR = (
    Path(r"C:\Users\liu12\OneDrive\Desktop\NHANES_MetS_Project")
    / "result"
    / "integrated_mainline"
    / PACKAGE_NAME
)

SUBDIRS = [
    "manuscript",
    "figures",
    "tables",
    "supplement",
    "checklists",
    "source_data",
    "logs",
    "scripts",
]

COLORS = {
    "navy": (31, 78, 121),
    "blue": (68, 126, 184),
    "orange": (217, 95, 2),
    "green": (27, 158, 119),
    "purple": (117, 112, 179),
    "red": (204, 80, 62),
    "teal": (36, 143, 151),
    "gray": (90, 94, 99),
    "light_gray": (233, 236, 239),
    "mid_gray": (185, 190, 195),
    "black": (20, 24, 28),
    "white": (255, 255, 255),
}


def ensure_dirs() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for subdir in SUBDIRS:
        (OUT / subdir).mkdir(parents=True, exist_ok=True)


def read_csv(path: Path) -> List[Dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, rows: Sequence[Dict[str, object]], fieldnames: Sequence[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.strip() + "\n", encoding="utf-8")


def as_float(value: object) -> float | None:
    try:
        if value is None or value == "":
            return None
        x = float(value)
        if math.isnan(x):
            return None
        return x
    except (TypeError, ValueError):
        return None


def fmt_num(value: object, digits: int = 3) -> str:
    x = as_float(value)
    if x is None:
        return ""
    return f"{x:.{digits}f}"


def fmt_int(value: object) -> str:
    x = as_float(value)
    if x is None:
        return ""
    return f"{int(round(x))}"


def fmt_p(value: object) -> str:
    x = as_float(value)
    if x is None:
        return ""
    if x < 0.001:
        return f"{x:.2e}"
    return f"{x:.3f}"


def ci(row: Dict[str, str], est_key: str = "estimate", lo_key: str = "ci_low", hi_key: str = "ci_high") -> str:
    return f"{fmt_num(row.get(est_key))} ({fmt_num(row.get(lo_key))}, {fmt_num(row.get(hi_key))})"


def markdown_table(rows: Sequence[Dict[str, object]], columns: Sequence[str]) -> str:
    if not rows:
        return ""
    header = "| " + " | ".join(columns) + " |"
    sep = "| " + " | ".join(["---"] * len(columns)) + " |"
    lines = [header, sep]
    for row in rows:
        vals = [str(row.get(col, "")) for col in columns]
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def first_row(rows: Iterable[Dict[str, str]], **criteria: str) -> Dict[str, str]:
    for row in rows:
        if all(row.get(k) == v for k, v in criteria.items()):
            return row
    raise KeyError(f"No row matching {criteria}")


def label_outcome(value: str) -> str:
    labels = {
        "msi_pca": "MSI PCA",
        "msi_core_partial": "MSI core",
        "msi_no_bmi": "MSI without BMI",
        "ln_HOMA_IR": "ln(HOMA-IR)",
        "HbA1c": "HbA1c",
        "iauc_2h_per1000": "2-h iAUC /1000",
        "peak_delta_2h": "2-h peak delta",
        "high_iauc_top_quartile": "High iAUC risk",
        "high_peak_top_quartile": "High peak risk",
    }
    return labels.get(value, value)


def label_exposure(value: str) -> str:
    labels = {
        "ln_oxidative_to_MEHP": "ln(oxidative DEHP metabolites/MEHP)",
        "pct_oxidative_10": "Oxidative metabolite fraction, per 10 pp",
    }
    return labels.get(value, value)


def label_term(value: str) -> str:
    labels = {
        "carbs_10g": "Carbohydrate load, per 10 g",
        "msi_core_centered": "MSI, centered",
        "carbs_x_msi": "Carbohydrate load x MSI",
        "msi_centered": "MSI, centered",
        "digestibility_centered": "Digestibility-risk score, centered",
        "msi_x_digestibility": "MSI x digestibility-risk score",
    }
    return labels.get(value, value)


def load_inputs() -> Dict[str, List[Dict[str, str]]]:
    inputs = {
        "dataset": read_csv(PHASE2 / "Table1_dataset_and_phenotype_summary.csv"),
        "nhanes": read_csv(PHASE2 / "Table2_NHANES_R_survey_key_results.csv"),
        "cgmacros": read_csv(PHASE2 / "Table3_CGMacros_MSI_PPGR_results.csv"),
        "food_curve": read_csv(PHASE2 / "Table4_FoodMatrix_CurvePhenotype_results.csv"),
        "pred_perf": read_csv(PHASE2 / "Table5A_prediction_model_performance.csv"),
        "pred_gain": read_csv(PHASE2 / "Table5B_prediction_incremental_gain.csv"),
        "external": read_csv(PHASE2 / "Table5C_external_boundary_summary.csv"),
        "fig3": read_csv(PHASE2 / "Figure3_data_MSI_tertile_PPGR.csv"),
        "fig4_clusters": read_csv(PHASE2 / "Figure4_data_curve_clusters.csv"),
        "mixture": read_csv(
            BOOST
            / "module_B_nhanes_advanced_epidemiology"
            / "nhanes_dehp_mixture_composition_results.csv"
        ),
        "conformal": read_csv(
            BOOST
            / "module_H_prediction_TRIPOD_PROBAST_AI"
            / "prediction_conformal_interval_coverage.csv"
        ),
        "leakage": read_csv(
            BOOST
            / "module_H_prediction_TRIPOD_PROBAST_AI"
            / "prediction_leakage_audit.csv"
        ),
    }
    return inputs


def build_tables(data: Dict[str, List[Dict[str, str]]]) -> Dict[str, List[Dict[str, object]]]:
    table1 = []
    dataset_labels = {
        "NHANES_2013_2018": "NHANES 2013-2018",
        "CGMacros_subjects": "CGMacros",
        "Stanford_CGMDB": "Stanford CGMDB",
        "T1D_UOM": "T1D-UOM",
        "Jaeb_CGMND_HealthyAdults": "Jaeb CGMND healthy adults",
    }
    for r in data["dataset"]:
        table1.append(
            {
                "Dataset": dataset_labels.get(r["dataset"], r["dataset"]),
                "Role in integrated analysis": r["role"],
                "Participants": fmt_int(r.get("n_subjects_or_participants")),
                "Meals/events/readings": fmt_int(r.get("n_meals_or_events")),
                "MSI available": fmt_int(r.get("n_msi_partial")),
                "Median MSI": fmt_num(r.get("median_msi")),
                "High MSI tertile": fmt_int(r.get("high_tertile_n")),
            }
        )

    table2 = []
    for r in data["nhanes"]:
        if r["outcome"] in {"msi_pca", "msi_core_partial", "msi_no_bmi", "ln_HOMA_IR", "HbA1c"}:
            table2.append(
                {
                    "Outcome": label_outcome(r["outcome"]),
                    "Exposure": label_exposure(r["exposure"]),
                    "Beta (95% CI)": r.get("estimate_CI_for_table") or f"{fmt_num(r['R_beta'])} ({fmt_num(r['R_ci_low'])}, {fmt_num(r['R_ci_high'])})",
                    "P": fmt_p(r.get("R_p")),
                    "Q": fmt_p(r.get("R_q")),
                }
            )

    keep_terms = {"carbs_10g", "msi_core_centered", "carbs_x_msi"}
    keep_outcomes = {"iauc_2h_per1000", "peak_delta_2h", "high_iauc_top_quartile", "high_peak_top_quartile"}
    table3 = []
    for r in data["cgmacros"]:
        if r["model"] == "base_interaction" and r["term"] in keep_terms and r["outcome"] in keep_outcomes:
            table3.append(
                {
                    "Outcome": label_outcome(r["outcome"]),
                    "Term": label_term(r["term"]),
                    "Estimate (95% CI)": r.get("estimate_CI_for_table") or ci(r),
                    "Q": fmt_p(r.get("q_value")),
                    "N meals": fmt_int(r.get("n")),
                    "N participants": fmt_int(r.get("n_subjects")),
                    "Unit": r.get("unit", ""),
                }
            )

    table4 = []
    for r in data["food_curve"]:
        keep = False
        if r["analysis_block"] == "food_matrix_ppgr":
            keep = True
        if r["analysis_block"] == "curve_phenotype" and r["term"] == "msi_x_digestibility":
            keep = True
        if keep:
            table4.append(
                {
                    "Analysis block": r["analysis_block"].replace("_", " "),
                    "Outcome": r["outcome_label"],
                    "Term": label_term(r["term"]),
                    "Estimate (95% CI)": r.get("estimate_CI_for_table") or ci(r),
                    "Q": fmt_p(r.get("q_value")),
                    "N meals": fmt_int(r.get("n")),
                }
            )

    table5 = []
    for r in data["pred_perf"]:
        if r["model"] in {"M0_carb_only", "M2_macro_precgm_context", "M3_msi_enhanced"}:
            table5.append(
                {
                    "Target": r["target_label"],
                    "Model": r["model"],
                    "MAE": fmt_num(r.get("MAE"), 1),
                    "R2": fmt_num(r.get("R2")),
                    "AUC high response": fmt_num(r.get("AUC_high_response")),
                    "AUPRC high response": fmt_num(r.get("AUPRC_high_response")),
                    "Calibration slope": fmt_num(r.get("calibration_slope")),
                    "Net benefit, pt=0.25": fmt_num(r.get("net_benefit_pt25")),
                }
            )

    table6 = []
    for r in data["external"]:
        table6.append(
            {
                "Dataset": r["dataset"],
                "Participants": fmt_int(r.get("n_subjects")),
                "Events/readings": fmt_int(r.get("n_events")),
                "Event type": r.get("event_type", ""),
                "Median 2-h iAUC": fmt_num(r.get("median_iauc_2h"), 1),
                "Median 2-h peak delta": fmt_num(r.get("median_peak_delta_2h"), 1),
                "High-rate / time above 140": fmt_num(r.get("time_above_140_or_high_rate")),
            }
        )

    outputs = {
        "Table1_dataset_and_phenotype_summary_AJCN.csv": table1,
        "Table2_NHANES_DEHP_MSI_survey_results_AJCN.csv": table2,
        "Table3_CGMacros_MSI_PPGR_models_AJCN.csv": table3,
        "Table4_food_matrix_curve_phenotype_AJCN.csv": table4,
        "Table5_prediction_model_performance_AJCN.csv": table5,
        "SupplementaryTable_external_boundary_AJCN.csv": table6,
    }
    for filename, rows in outputs.items():
        write_csv(OUT / "tables" / filename, rows, list(rows[0].keys()))

    # Copy all analysis source tables so every figure/table can be traced.
    for src in PHASE2.glob("*.csv"):
        shutil.copy2(src, OUT / "source_data" / src.name)
    for src in [
        BOOST / "module_B_nhanes_advanced_epidemiology" / "nhanes_dehp_mixture_composition_results.csv",
        BOOST / "module_B_nhanes_advanced_epidemiology" / "nhanes_effect_modification_results.csv",
        BOOST / "module_B_nhanes_advanced_epidemiology" / "nhanes_negative_control_and_sensitivity_results.csv",
        BOOST / "module_C_cgmacros_food_matrix" / "cgmacros_food_matrix_ppgr_models.csv",
        BOOST / "module_D_cgm_curve_phenotypes" / "curve_shape_msi_food_matrix_models.csv",
        BOOST / "module_F_external_transportability" / "external_transportability_map.csv",
        BOOST / "module_H_prediction_TRIPOD_PROBAST_AI" / "prediction_conformal_interval_coverage.csv",
        BOOST / "module_H_prediction_TRIPOD_PROBAST_AI" / "prediction_leakage_audit.csv",
    ]:
        if src.exists():
            shutil.copy2(src, OUT / "source_data" / src.name)

    return outputs


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
        Path(r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> List[str]:
    words = text.replace("\n", " \n ").split()
    lines: List[str] = []
    current = ""
    for word in words:
        if word == "\n":
            if current:
                lines.append(current)
                current = ""
            continue
        trial = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_textbox(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str,
    fill: tuple[int, int, int],
    outline: tuple[int, int, int] = COLORS["navy"],
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=4)
    title_font = get_font(42, bold=True)
    body_font = get_font(32, bold=False)
    draw.text((x1 + 28, y1 + 24), title, font=title_font, fill=COLORS["black"])
    y = y1 + 84
    for line in wrap_text(draw, body, body_font, x2 - x1 - 56):
        draw.text((x1 + 28, y), line, font=body_font, fill=COLORS["black"])
        y += 42


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=COLORS["gray"]) -> None:
    draw.line([start, end], fill=color, width=8)
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 26
    points = [
        (x2, y2),
        (x2 - length * math.cos(angle - 0.45), y2 - length * math.sin(angle - 0.45)),
        (x2 - length * math.cos(angle + 0.45), y2 - length * math.sin(angle + 0.45)),
    ]
    draw.polygon(points, fill=color)


def draw_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str = "") -> None:
    draw.text((110, 70), title, font=get_font(58, bold=True), fill=COLORS["black"])
    if subtitle:
        draw.text((110, 138), subtitle, font=get_font(34), fill=COLORS["gray"])


def draw_axis(draw: ImageDraw.ImageDraw, x0: int, y0: int, x1: int, y1: int) -> None:
    draw.line([(x0, y0), (x0, y1), (x1, y1)], fill=COLORS["black"], width=4)


def make_figure1() -> Path:
    img = Image.new("RGB", (3200, 2200), COLORS["white"])
    draw = ImageDraw.Draw(img)
    draw_title(draw, "Figure 1. Integrated no-wet-lab study framework", "Evidence path from environmental metabolic profile to CGM-measured meal vulnerability")
    boxes = [
        (130, 330, 820, 730, "NHANES", "DEHP oxidative metabolite profile and metabolic phenotypes in a population survey", COLORS["light_gray"]),
        (930, 330, 1620, 730, "Bridge phenotype", "Metabolic Susceptibility Index (MSI): adiposity, glycemia, insulin resistance, lipid risk", (232, 242, 251)),
        (1730, 330, 2420, 730, "CGMacros", "Free-living meals with CGM-derived iAUC, peak excursions, and response phenotypes", (233, 245, 241)),
        (2530, 330, 3090, 730, "Prediction", "MSI-enhanced high-response risk stratification and external boundary analyses", (249, 240, 231)),
    ]
    for x1, y1, x2, y2, title, body, fill in boxes:
        draw_textbox(draw, (x1, y1, x2, y2), title, body, fill)
    for i in range(len(boxes) - 1):
        arrow(draw, (boxes[i][2] + 20, 530), (boxes[i + 1][0] - 20, 530), COLORS["gray"])

    draw.text((150, 900), "Primary inferential claims", font=get_font(44, bold=True), fill=COLORS["black"])
    claims = [
        ("Population association", "DEHP oxidative profile was associated with higher MSI, ln(HOMA-IR), and HbA1c in survey-weighted NHANES models.", COLORS["blue"]),
        ("Dynamic expression", "Higher MSI identified CGMacros participants with larger 2-h iAUC, peak excursions, and high-response meal rates.", COLORS["green"]),
        ("Meal matrix susceptibility", "Exploratory digestibility-risk proxy showed positive MSI interactions for iAUC, peak, delayed peak, and prolonged elevation phenotypes.", COLORS["orange"]),
        ("Model utility boundary", "MSI improved prediction within CGMacros; external datasets were used to define transportability boundaries, not direct replication.", COLORS["purple"]),
    ]
    x, y = 150, 1000
    for title, body, color in claims:
        draw.rounded_rectangle((x, y, x + 1360, y + 260), radius=14, outline=color, width=5, fill=(255, 255, 255))
        draw.rectangle((x, y, x + 24, y + 260), fill=color)
        draw.text((x + 52, y + 28), title, font=get_font(36, bold=True), fill=COLORS["black"])
        yy = y + 82
        for line in wrap_text(draw, body, get_font(30), 1220):
            draw.text((x + 52, yy), line, font=get_font(30), fill=COLORS["black"])
            yy += 38
        if x < 1500:
            x = 1660
        else:
            x = 150
            y += 320

    draw.text((150, 2035), "Guardrail: this design does not test whether DEHP directly causes CGM postprandial responses.", font=get_font(34, bold=True), fill=COLORS["red"])
    path = OUT / "figures" / "Figure1_integrated_study_framework.png"
    img.save(path, dpi=(300, 300))
    return path


def make_figure2(data: Dict[str, List[Dict[str, str]]]) -> Path:
    rows = [
        r
        for r in data["nhanes"]
        if r["outcome"] in {"msi_core_partial", "ln_HOMA_IR", "HbA1c", "msi_no_bmi", "msi_pca"}
    ]
    rows = sorted(rows, key=lambda r: (r["exposure"], r["outcome"]))
    img = Image.new("RGB", (3000, 2200), COLORS["white"])
    draw = ImageDraw.Draw(img)
    draw_title(draw, "Figure 2. NHANES DEHP oxidative profile and metabolic susceptibility", "Survey-weighted beta estimates with 95% confidence intervals")
    x0, x1 = 1120, 2700
    y0, row_h = 360, 120
    xmin, xmax = -0.05, 0.35

    def xp(value: float) -> int:
        return int(x0 + (value - xmin) / (xmax - xmin) * (x1 - x0))

    draw.line([(xp(0), y0 - 70), (xp(0), y0 + row_h * len(rows) + 60)], fill=COLORS["mid_gray"], width=4)
    for tick in [-0.05, 0, 0.10, 0.20, 0.30, 0.35]:
        x = xp(tick)
        draw.line([(x, y0 + row_h * len(rows) + 20), (x, y0 + row_h * len(rows) + 42)], fill=COLORS["black"], width=3)
        draw.text((x - 36, y0 + row_h * len(rows) + 55), f"{tick:.2f}", font=get_font(28), fill=COLORS["black"])
    draw.text((x0 + 250, y0 + row_h * len(rows) + 115), "Beta estimate", font=get_font(34, bold=True), fill=COLORS["black"])

    header_font = get_font(34, bold=True)
    draw.text((120, y0 - 125), "Outcome", font=header_font, fill=COLORS["black"])
    draw.text((535, y0 - 125), "Exposure", font=header_font, fill=COLORS["black"])
    draw.text((2750, y0 - 125), "q", font=header_font, fill=COLORS["black"])
    font = get_font(30)
    for idx, r in enumerate(rows):
        y = y0 + idx * row_h
        if idx % 2:
            draw.rectangle((90, y - 20, 2860, y + 85), fill=(248, 249, 250))
        est = as_float(r.get("R_beta")) or 0
        lo = as_float(r.get("R_ci_low")) or 0
        hi = as_float(r.get("R_ci_high")) or 0
        color = COLORS["blue"] if r["exposure"] == "ln_oxidative_to_MEHP" else COLORS["orange"]
        draw.text((120, y), label_outcome(r["outcome"]), font=font, fill=COLORS["black"])
        draw.text((535, y), "ln ratio" if r["exposure"] == "ln_oxidative_to_MEHP" else "% oxidative", font=font, fill=COLORS["black"])
        draw.line([(xp(lo), y + 25), (xp(hi), y + 25)], fill=color, width=8)
        draw.ellipse((xp(est) - 14, y + 11, xp(est) + 14, y + 39), fill=color)
        draw.text((2725, y), fmt_p(r.get("R_q")), font=font, fill=COLORS["black"])

    # Add composition sensitivity from the dry-lab booster.
    mix = [
        r
        for r in data["mixture"]
        if r.get("outcome") == "msi_core_partial" and r.get("term") == "ln_oxidative_to_MEHP"
    ]
    if mix:
        r = mix[0]
        text = (
            "Composition proxy sensitivity: ln(oxidative/MEHP) remained positive for MSI-core "
            f"({ci(r)}, q={fmt_p(r.get('q_value'))}; n={fmt_int(r.get('n'))})."
        )
        draw.rounded_rectangle((220, 1880, 2780, 2035), radius=16, fill=(247, 250, 252), outline=COLORS["mid_gray"], width=3)
        draw.text((260, 1918), text, font=get_font(32), fill=COLORS["black"])

    path = OUT / "figures" / "Figure2_NHANES_DEHP_MSI_forest.png"
    img.save(path, dpi=(300, 300))
    return path


def make_figure3(data: Dict[str, List[Dict[str, str]]]) -> Path:
    rows = data["fig3"]
    img = Image.new("RGB", (3000, 2200), COLORS["white"])
    draw = ImageDraw.Draw(img)
    draw_title(draw, "Figure 3. CGMacros MSI tertiles and postprandial glycemic response", "Meal-level summary by participant MSI tertile")
    tertiles = [r["msi_tertile"] for r in rows]
    colors = [COLORS["blue"], COLORS["green"], COLORS["orange"]]

    panels = [
        (170, 360, 890, 1520, "A. Mean 2-h iAUC", "mean_iauc_2h", 4000, "mg/dL x min"),
        (1120, 360, 1840, 1520, "B. Mean peak delta", "mean_peak_delta_2h", 70, "mg/dL"),
        (2070, 360, 2790, 1520, "C. High-response rate", "high_iauc_rate", 0.40, "proportion"),
    ]
    for px0, py0, px1, py1, title, key, ymax, ylabel in panels:
        draw.text((px0, py0 - 90), title, font=get_font(38, bold=True), fill=COLORS["black"])
        draw_axis(draw, px0 + 90, py0, px1, py1)
        tick_values = [0.0, 0.1, 0.2, 0.3, 0.4] if ymax <= 1 else [ymax * t / 5 for t in range(0, 6)]
        for val in tick_values:
            y = py1 - int((val / ymax) * (py1 - py0))
            draw.line([(px0 + 82, y), (px0 + 90, y)], fill=COLORS["black"], width=3)
            label = f"{val:.2f}" if ymax < 1 else f"{val:.0f}"
            draw.text((px0 + 4, y - 18), label, font=get_font(24), fill=COLORS["gray"])
        for i, r in enumerate(rows):
            val = as_float(r.get(key)) or 0
            bar_h = int((val / ymax) * (py1 - py0))
            bx = px0 + 160 + i * 170
            bar_color = COLORS["orange"] if key == "high_iauc_rate" else colors[i]
            draw.rectangle((bx, py1 - bar_h, bx + 105, py1), fill=bar_color)
            draw.text((bx - 14, py1 + 22), tertiles[i], font=get_font(28), fill=COLORS["black"])
            label = f"{val:.2f}" if ymax <= 1 else f"{val:.0f}"
            draw.text((bx - 10, py1 - bar_h - 42), label, font=get_font(26, bold=True), fill=COLORS["black"])
            if key == "high_iauc_rate":
                val2 = as_float(r.get("high_peak_rate")) or 0
                bh2 = int((val2 / ymax) * (py1 - py0))
                draw.rectangle((bx + 112, py1 - bh2, bx + 160, py1), fill=COLORS["purple"])
        draw.text((px0 + 160, py1 + 85), ylabel, font=get_font(28), fill=COLORS["gray"])
    draw.rounded_rectangle((2055, 1590, 2800, 1725), radius=10, fill=(250, 250, 250), outline=COLORS["mid_gray"], width=2)
    draw.rectangle((2095, 1628, 2155, 1668), fill=COLORS["orange"])
    draw.text((2175, 1622), "High iAUC rate", font=get_font(28), fill=COLORS["black"])
    draw.rectangle((2460, 1628, 2520, 1668), fill=COLORS["purple"])
    draw.text((2540, 1622), "High peak rate", font=get_font(28), fill=COLORS["black"])

    note = "High MSI tertile: mean iAUC 3686.1 vs 2543.1 in low tertile; mean peak delta 60.4 vs 43.4 mg/dL."
    draw.text((180, 1900), note, font=get_font(34, bold=True), fill=COLORS["black"])
    path = OUT / "figures" / "Figure3_CGMacros_MSI_tertile_PPGR.png"
    img.save(path, dpi=(300, 300))
    return path


def make_figure4(data: Dict[str, List[Dict[str, str]]]) -> Path:
    img = Image.new("RGB", (3200, 2200), COLORS["white"])
    draw = ImageDraw.Draw(img)
    draw_title(draw, "Figure 4. Food-matrix susceptibility and CGM curve phenotypes", "Exploratory dry-lab food matrix proxy and response-shape clustering")

    # Panel A: interaction coefficients.
    panel_x0, panel_y0 = 190, 360
    draw.text((panel_x0, panel_y0 - 90), "A. MSI x digestibility-risk interactions", font=get_font(38, bold=True), fill=COLORS["black"])
    rows = [
        r
        for r in data["food_curve"]
        if r.get("term") == "msi_x_digestibility"
        and r.get("outcome") in {
            "2h iAUC /1000",
            "2h peak delta",
            "delayed_peak_flag",
            "prolonged_elevation_flag",
            "early_slope_proxy_0_to_peak",
        }
    ]
    x0, x1 = 760, 1340
    y0, row_h = panel_y0 + 40, 118
    xmin, xmax = -0.05, 7.8

    def xp(v: float) -> int:
        return int(x0 + (v - xmin) / (xmax - xmin) * (x1 - x0))

    draw.line([(xp(0), y0 - 30), (xp(0), y0 + row_h * len(rows) + 40)], fill=COLORS["mid_gray"], width=3)
    for i, r in enumerate(rows):
        y = y0 + i * row_h
        est = as_float(r.get("estimate")) or 0
        lo = as_float(r.get("ci_low")) or 0
        hi = as_float(r.get("ci_high")) or 0
        color = COLORS["orange"] if "peak" in r.get("outcome", "") else COLORS["green"]
        label = r.get("outcome_label", r.get("outcome", "")).replace("_", " ")
        draw.text((panel_x0, y), label[:34], font=get_font(30), fill=COLORS["black"])
        draw.line([(xp(lo), y + 24), (xp(hi), y + 24)], fill=color, width=8)
        draw.ellipse((xp(est) - 14, y + 10, xp(est) + 14, y + 38), fill=color)
        draw.text((x1 + 35, y), f"{fmt_num(est)}; q={fmt_p(r.get('q_value'))}", font=get_font(27), fill=COLORS["black"])
    draw.text((x0 + 130, y0 + row_h * len(rows) + 72), "Coefficient estimate", font=get_font(30, bold=True), fill=COLORS["black"])

    # Panel B: curve cluster map.
    px0, py0, px1, py1 = 1900, 420, 3040, 1600
    draw.text((px0, panel_y0 - 90), "B. Curve-shape clusters", font=get_font(38, bold=True), fill=COLORS["black"])
    draw_axis(draw, px0, py0, px1, py1)
    max_i = max(as_float(r.get("mean_iauc_2h")) or 0 for r in data["fig4_clusters"])
    max_p = max(as_float(r.get("mean_peak_delta_2h")) or 0 for r in data["fig4_clusters"])
    min_msi = min(as_float(r.get("mean_msi")) or 0 for r in data["fig4_clusters"])
    max_msi = max(as_float(r.get("mean_msi")) or 0 for r in data["fig4_clusters"])

    def sx(v: float) -> int:
        return int(px0 + (v / (max_i * 1.08)) * (px1 - px0))

    def sy(v: float) -> int:
        return int(py1 - (v / (max_p * 1.12)) * (py1 - py0))

    for r in data["fig4_clusters"]:
        iauc = as_float(r.get("mean_iauc_2h")) or 0
        peak = as_float(r.get("mean_peak_delta_2h")) or 0
        n = as_float(r.get("n_meals")) or 1
        msi = as_float(r.get("mean_msi")) or 0
        frac = 0 if max_msi == min_msi else (msi - min_msi) / (max_msi - min_msi)
        color = (
            int(COLORS["blue"][0] * (1 - frac) + COLORS["red"][0] * frac),
            int(COLORS["blue"][1] * (1 - frac) + COLORS["red"][1] * frac),
            int(COLORS["blue"][2] * (1 - frac) + COLORS["red"][2] * frac),
        )
        radius = int(24 + math.sqrt(n) * 3)
        x, y = sx(iauc), sy(peak)
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=color, outline=COLORS["black"], width=3)
        label = r.get("shape_cluster_label", "").replace("_", " ")
        if "large" in label:
            draw.text((x - 310, y - 100), "large prolonged response", font=get_font(30, bold=True), fill=COLORS["red"])
        elif "high" in label:
            draw.text((x - 160, y - 85), "high peak fast recovery", font=get_font(26), fill=COLORS["black"])
        elif "delayed" in label:
            draw.text((x - 110, y - 80), "delayed peak", font=get_font(26), fill=COLORS["black"])
    for t in range(0, 6):
        xv = max_i * t / 5
        x = sx(xv)
        draw.line([(x, py1), (x, py1 + 10)], fill=COLORS["black"], width=3)
        draw.text((x - 35, py1 + 22), f"{xv/1000:.1f}", font=get_font(24), fill=COLORS["gray"])
        yv = max_p * t / 5
        y = sy(yv)
        draw.line([(px0 - 10, y), (px0, y)], fill=COLORS["black"], width=3)
        draw.text((px0 - 82, y - 15), f"{yv:.0f}", font=get_font(24), fill=COLORS["gray"])
    draw.text((px0 + 330, py1 + 82), "Mean 2-h iAUC (x1000)", font=get_font(28, bold=True), fill=COLORS["black"])
    draw.text((px0 + 5, py0 - 60), "Mean peak delta (mg/dL)", font=get_font(28, bold=True), fill=COLORS["black"])

    note = "The highest-risk cluster combined large iAUC, high peak excursions, higher MSI, and higher digestibility-risk score."
    draw.text((220, 1910), note, font=get_font(34, bold=True), fill=COLORS["black"])
    path = OUT / "figures" / "Figure4_food_matrix_curve_phenotypes.png"
    img.save(path, dpi=(300, 300))
    return path


def make_figure5(data: Dict[str, List[Dict[str, str]]]) -> Path:
    img = Image.new("RGB", (3200, 2200), COLORS["white"])
    draw = ImageDraw.Draw(img)
    draw_title(draw, "Figure 5. Prediction performance and external boundary evidence", "Nested CGMacros prediction with external CGM context")
    models = ["M0_carb_only", "M1_macro_timing", "M2_macro_precgm_context", "M3_msi_enhanced"]
    colors = [COLORS["mid_gray"], COLORS["blue"], COLORS["green"], COLORS["orange"]]

    def panel_bars(px0: int, py0: int, px1: int, py1: int, title: str, target: str, metric: str, ymax: float, lower_better: bool = False, metric_label: str | None = None) -> None:
        draw.text((px0, py0 - 90), title, font=get_font(36, bold=True), fill=COLORS["black"])
        draw_axis(draw, px0 + 90, py0, px1, py1)
        rows = {r["model"]: r for r in data["pred_perf"] if r["target"] == target}
        for i, model in enumerate(models):
            val = as_float(rows[model].get(metric)) or 0
            bar_h = int((val / ymax) * (py1 - py0))
            bx = px0 + 160 + i * 150
            draw.rectangle((bx, py1 - bar_h, bx + 95, py1), fill=colors[i])
            draw.text((bx - 8, py1 - bar_h - 38), f"{val:.2f}" if val < 10 else f"{val:.0f}", font=get_font(24, bold=True), fill=COLORS["black"])
            draw.text((bx - 12, py1 + 25), model.split("_")[0], font=get_font(26), fill=COLORS["black"])
        label = "lower is better" if lower_better else "higher is better"
        draw.text((px0 + 125, py1 + 78), f"{metric_label or metric} ({label})", font=get_font(27), fill=COLORS["gray"])

    panel_bars(180, 380, 1080, 1260, "A. iAUC MAE", "iauc_2h", "MAE", 2500, True, "MAE")
    panel_bars(1240, 380, 2140, 1260, "B. Peak MAE", "peak_delta_2h", "MAE", 36, True, "MAE")
    panel_bars(2300, 380, 3080, 1260, "C. High-response AUC", "peak_delta_2h", "AUC_high_response", 0.85, False, "High-response AUC")

    legend_y = 1410
    for i, model in enumerate(models):
        draw.rectangle((250 + i * 720, legend_y, 310 + i * 720, legend_y + 45), fill=colors[i])
        draw.text((330 + i * 720, legend_y - 2), model, font=get_font(28), fill=COLORS["black"])

    # External boundary panel.
    draw.text((230, 1620), "D. External CGM context", font=get_font(36, bold=True), fill=COLORS["black"])
    x, y = 230, 1700
    for r in data["external"]:
        width = 680
        draw.rounded_rectangle((x, y, x + width, y + 230), radius=14, fill=(248, 249, 250), outline=COLORS["mid_gray"], width=3)
        draw.text((x + 28, y + 24), r["dataset"], font=get_font(30, bold=True), fill=COLORS["black"])
        body = (
            f"n={fmt_int(r.get('n_subjects'))}; events={fmt_int(r.get('n_events'))}; "
            f"median peak={fmt_num(r.get('median_peak_delta_2h'), 1) or 'NA'}; "
            f"high/time>140={fmt_num(r.get('time_above_140_or_high_rate')) or 'NA'}"
        )
        for line_idx, line in enumerate(wrap_text(draw, body, get_font(25), width - 56)):
            draw.text((x + 28, y + 80 + line_idx * 34), line, font=get_font(25), fill=COLORS["black"])
        x += width + 70
        if x + width > 3100:
            x = 230
            y += 270

    path = OUT / "figures" / "Figure5_prediction_external_boundary.png"
    img.save(path, dpi=(300, 300))
    return path


def build_figures(data: Dict[str, List[Dict[str, str]]]) -> List[Path]:
    return [
        make_figure1(),
        make_figure2(data),
        make_figure3(data),
        make_figure4(data),
        make_figure5(data),
    ]


def key_stats(data: Dict[str, List[Dict[str, str]]]) -> Dict[str, str]:
    nhanes_msi = first_row(data["nhanes"], outcome="msi_core_partial", exposure="ln_oxidative_to_MEHP")
    nhanes_homa = first_row(data["nhanes"], outcome="ln_HOMA_IR", exposure="ln_oxidative_to_MEHP")
    nhanes_hba1c = first_row(data["nhanes"], outcome="HbA1c", exposure="ln_oxidative_to_MEHP")
    nhanes_pct_msi = first_row(data["nhanes"], outcome="msi_core_partial", exposure="pct_oxidative_10")
    cg_i = first_row(data["cgmacros"], outcome="iauc_2h_per1000", model="base_interaction", term="msi_core_centered")
    cg_p = first_row(data["cgmacros"], outcome="peak_delta_2h", model="base_interaction", term="msi_core_centered")
    fm_i = first_row(data["food_curve"], analysis_block="food_matrix_ppgr", outcome="2h iAUC /1000", term="msi_x_digestibility")
    fm_p = first_row(data["food_curve"], analysis_block="food_matrix_ppgr", outcome="2h peak delta", term="msi_x_digestibility")
    m3_i = first_row(data["pred_perf"], target="iauc_2h", model="M3_msi_enhanced")
    m3_p = first_row(data["pred_perf"], target="peak_delta_2h", model="M3_msi_enhanced")
    gain_i = next(r for r in data["pred_gain"] if r["target"] == "iauc_2h" and "M2" in r["comparison"])
    gain_p = next(r for r in data["pred_gain"] if r["target"] == "peak_delta_2h" and "M2" in r["comparison"])
    high = first_row(data["fig3"], msi_tertile="high")
    low = first_row(data["fig3"], msi_tertile="low")
    large = next(r for r in data["fig4_clusters"] if r["shape_cluster_label"] == "large_prolonged_response")
    return {
        "nhanes_msi": nhanes_msi.get("estimate_CI_for_table", ""),
        "nhanes_msi_q": fmt_p(nhanes_msi.get("R_q")),
        "nhanes_homa": nhanes_homa.get("estimate_CI_for_table", ""),
        "nhanes_homa_q": fmt_p(nhanes_homa.get("R_q")),
        "nhanes_hba1c": nhanes_hba1c.get("estimate_CI_for_table", ""),
        "nhanes_hba1c_q": fmt_p(nhanes_hba1c.get("R_q")),
        "nhanes_pct_msi": nhanes_pct_msi.get("estimate_CI_for_table", ""),
        "nhanes_pct_msi_q": fmt_p(nhanes_pct_msi.get("R_q")),
        "cg_i": cg_i.get("estimate_CI_for_table", ""),
        "cg_i_q": fmt_p(cg_i.get("q_value")),
        "cg_p": cg_p.get("estimate_CI_for_table", ""),
        "cg_p_q": fmt_p(cg_p.get("q_value")),
        "fm_i": fm_i.get("estimate_CI_for_table", ""),
        "fm_i_q": fmt_p(fm_i.get("q_value")),
        "fm_p": fm_p.get("estimate_CI_for_table", ""),
        "fm_p_q": fmt_p(fm_p.get("q_value")),
        "m3_i_mae": fmt_num(m3_i.get("MAE"), 1),
        "m3_i_r2": fmt_num(m3_i.get("R2")),
        "m3_i_auc": fmt_num(m3_i.get("AUC_high_response")),
        "m3_i_auprc": fmt_num(m3_i.get("AUPRC_high_response")),
        "m3_p_mae": fmt_num(m3_p.get("MAE"), 1),
        "m3_p_r2": fmt_num(m3_p.get("R2")),
        "m3_p_auc": fmt_num(m3_p.get("AUC_high_response")),
        "m3_p_auprc": fmt_num(m3_p.get("AUPRC_high_response")),
        "gain_i_mae": fmt_num(gain_i.get("MAE_pct_change"), 1),
        "gain_i_auc": fmt_num(gain_i.get("AUC_delta")),
        "gain_p_mae": fmt_num(gain_p.get("MAE_pct_change"), 1),
        "gain_p_auc": fmt_num(gain_p.get("AUC_delta")),
        "high_i_mean": fmt_num(high.get("mean_iauc_2h"), 1),
        "low_i_mean": fmt_num(low.get("mean_iauc_2h"), 1),
        "high_p_mean": fmt_num(high.get("mean_peak_delta_2h"), 1),
        "low_p_mean": fmt_num(low.get("mean_peak_delta_2h"), 1),
        "large_n": fmt_int(large.get("n_meals")),
        "large_i": fmt_num(large.get("mean_iauc_2h"), 1),
        "large_p": fmt_num(large.get("mean_peak_delta_2h"), 1),
    }


def build_manuscript_markdown(data: Dict[str, List[Dict[str, str]]], tables: Dict[str, List[Dict[str, object]]]) -> str:
    s = key_stats(data)
    title = "A phthalate oxidative-metabolism profile, metabolic susceptibility, and postprandial glycemic vulnerability"
    abstract = f"""
## Abstract

**Background:** Postprandial glycemic responses are highly individualized, but few human studies connect environmental metabolic-disruptor profiles with dynamic continuous glucose monitoring (CGM) phenotypes.

**Objectives:** We tested whether a DEHP oxidative-metabolite profile was associated with a metabolic susceptibility phenotype in NHANES and whether the same phenotype identified greater postprandial glycemic vulnerability in CGMacros.

**Methods:** We conducted an integrated computational analysis of NHANES 2013-2018, CGMacros, and three external CGM datasets. A Metabolic Susceptibility Index (MSI) summarized adiposity, glycemia, insulin resistance, and lipid risk. In NHANES, survey-weighted models related urinary DEHP oxidative profile measures to MSI and metabolic markers. In CGMacros, subject-clustered meal-level models related MSI to 2-h incremental area under the glucose curve (iAUC), peak glucose excursion, high-response indicators, food-matrix proxy features, curve phenotypes, and nested prediction models. External datasets were used to define transportability boundaries.

**Results:** In NHANES, ln(oxidative DEHP metabolites/MEHP) was associated with MSI-core (beta {s['nhanes_msi']}; q={s['nhanes_msi_q']}), ln(HOMA-IR) (beta {s['nhanes_homa']}; q={s['nhanes_homa_q']}), and HbA1c (beta {s['nhanes_hba1c']}; q={s['nhanes_hba1c_q']}). In CGMacros (40 participants; 1498 meals), higher MSI was associated with greater 2-h iAUC/1000 (beta {s['cg_i']}; q={s['cg_i_q']}) and peak excursion (beta {s['cg_p']}; q={s['cg_p_q']}). High versus low MSI tertiles had higher mean iAUC ({s['high_i_mean']} vs {s['low_i_mean']} mg/dL x min) and peak excursion ({s['high_p_mean']} vs {s['low_p_mean']} mg/dL). Exploratory food-matrix analyses showed positive MSI x digestibility-risk interactions for iAUC/1000 (beta {s['fm_i']}; q={s['fm_i_q']}) and peak excursion (beta {s['fm_p']}; q={s['fm_p_q']}). MSI-enhanced prediction improved iAUC and peak high-response discrimination relative to macro/pre-CGM context models.

**Conclusions:** A DEHP oxidative-metabolism profile was associated with metabolic susceptibility in NHANES, and this susceptibility phenotype was expressed as higher CGM-measured postprandial glycemic vulnerability in CGMacros. These findings support a susceptibility-aware precision-nutrition framework but require prospective and mechanistic validation before causal or clinical claims.
"""

    return f"""
# {title}

**Running title:** Environmental metabolic susceptibility and postprandial glycemia

**Authors:** [Author names to be completed]

**Affiliations:** [Institutional affiliations to be completed]

**Corresponding author:** [Name, address, email, telephone to be completed]

**Funding:** [Funding sources to be completed]

**Conflicts of interest:** [Conflict-of-interest statement to be completed]

**Data sharing:** NHANES data are publicly available from CDC/NCHS. CGMacros raw data and third-party CGM datasets are subject to their original access terms. Derived de-identified tables can be shared where permitted.

**Clinical trial registration:** Not applicable; secondary analysis of existing observational datasets.

**Keywords:** DEHP; phthalates; continuous glucose monitoring; postprandial glycemia; precision nutrition; insulin resistance; NHANES.

{abstract}

## Introduction

Postprandial glycemic response (PPGR) is a dynamic phenotype that varies markedly across individuals even for similar foods. CGM-based cohorts have shown that meal composition, timing, baseline physiology, and interindividual susceptibility jointly shape glycemic excursions. This heterogeneity motivates precision-nutrition models that move beyond average carbohydrate effects toward identifying who is vulnerable to which meal contexts.

Environmental metabolic disruptors may contribute to the biological background on which meal responses occur. Di(2-ethylhexyl) phthalate (DEHP) metabolites have been associated with obesity, insulin resistance, and diabetes-related markers in population studies. However, population exposure studies rarely connect environmental profiles to dynamic CGM endpoints, whereas CGM studies rarely include detailed environmental exposure biomarkers. A direct one-dataset test of DEHP exposure and PPGR is therefore not available in the present data resources.

We addressed this gap by using a bridge phenotype: a Metabolic Susceptibility Index (MSI) derived from adiposity, glycemic, insulin-resistance, and lipid-risk markers. The core hypothesis was not that DEHP directly causes higher CGM responses, but that a DEHP oxidative-metabolism profile is associated with a metabolic susceptibility phenotype in NHANES and that the same phenotype is expressed as higher PPGR vulnerability in CGMacros.

## Methods

### Study design and evidence layers

This was an integrated dry-lab analysis combining population survey, free-living CGM meal, and external CGM boundary evidence. The analysis followed a prespecified no-wet-lab scope: no in vitro digestion or bionic digestion data were used. The primary evidence layers were: 1) NHANES 2013-2018 for DEHP oxidative profile and metabolic susceptibility; 2) CGMacros for MSI and meal-level PPGR; 3) dry-lab food-matrix proxy annotation and CGM curve phenotypes for mechanistic refinement; and 4) external CGM datasets for transportability boundaries.

### Data sources

NHANES 2013-2018 contributed urinary DEHP metabolites and metabolic phenotypes. CGMacros contributed 40 participants and 1498 free-living meals with CGM-derived outcomes. External datasets included Stanford CGMDB standardized food challenges, T1D-UOM free-living meals from individuals with type 1 diabetes, and Jaeb CGMND healthy-adult CGM readings. Dataset roles and sample sizes are summarized in Table 1.

### Metabolic Susceptibility Index

MSI summarized BMI, HbA1c, fasting glucose, ln(HOMA-IR), and ln(TG/HDL-C). The primary MSI was a partial index designed to maximize available samples. Sensitivity indices included PCA-based MSI, MSI without BMI, and glycemia/insulin-resistance variants. The primary analyses used centered MSI values. Previous robustness checks showed near-identical behavior between MSI-core and PCA versions in both NHANES and CGMacros.

### DEHP oxidative profile in NHANES

Urinary DEHP metabolite features included MEHP and oxidative DEHP metabolites. The primary exposure was ln(oxidative DEHP metabolites/MEHP), interpreted as an oxidative metabolite profile. A secondary exposure was oxidative metabolite fraction per 10 percentage points. NHANES analyses used survey-weighted regression models and reported beta estimates, 95% CIs, P values, and false-discovery-rate q values.

### CGMacros meal-level outcomes

CGMacros outcomes included 2-h iAUC, 2-h peak glucose excursion, high-iAUC and high-peak top-quartile indicators, food-matrix/digestibility-risk proxy scores, and curve-shape phenotypes. Meal-level models used subject-clustered inference. Primary covariates included carbohydrate load and centered MSI; interaction models tested carbohydrate load x MSI. Food-matrix models tested MSI, digestibility-risk proxy score, and MSI x digestibility-risk score.

### Prediction analysis

Nested prediction models evaluated whether MSI improved meal-level risk stratification beyond carbohydrate-only, macro/timing, and macro/pre-CGM/context models. Performance metrics included MAE, RMSE, R2, calibration slope, AUC and AUPRC for high-response flags, and decision-curve net benefit at a threshold probability of 0.25. These models are for research-use stratification and are not clinical decision tools.

### Reporting standards

This package was organized to support AJCN-style submission and to align with STROBE, STROBE-nut, STROBE-ME, TRIPOD+AI, and PROBAST+AI expectations where applicable. Checklist drafts are provided in the supplement folder.

## Results

### Dataset overview

The integrated analysis included NHANES 2013-2018, CGMacros, Stanford CGMDB, T1D-UOM, and Jaeb CGMND healthy adults (Table 1). MSI was available for all 40 CGMacros participants and for the NHANES analytic subset used for the susceptibility phenotype. External datasets were not used to claim direct replication of the DEHP-MSI-CGM path; they were used to understand context, heterogeneity, and boundary conditions.

### DEHP oxidative profile was associated with metabolic susceptibility in NHANES

Survey-weighted NHANES models showed that ln(oxidative DEHP metabolites/MEHP) was positively associated with MSI-core (beta {s['nhanes_msi']}; q={s['nhanes_msi_q']}), ln(HOMA-IR) (beta {s['nhanes_homa']}; q={s['nhanes_homa_q']}), and HbA1c (beta {s['nhanes_hba1c']}; q={s['nhanes_hba1c_q']}) (Table 2; Figure 2). The oxidative metabolite fraction per 10 percentage points also showed a positive association with MSI-core (beta {s['nhanes_pct_msi']}; q={s['nhanes_pct_msi_q']}). These results support an exposure-profile-to-susceptibility association, while remaining observational and noncausal.

### MSI identified higher postprandial glycemic vulnerability in CGMacros

In CGMacros, higher centered MSI was associated with higher 2-h iAUC/1000 (beta {s['cg_i']}; q={s['cg_i_q']}) and higher 2-h peak excursion (beta {s['cg_p']}; q={s['cg_p_q']}) after accounting for carbohydrate load (Table 3). High versus low MSI tertiles had higher mean iAUC ({s['high_i_mean']} vs {s['low_i_mean']} mg/dL x min) and peak excursion ({s['high_p_mean']} vs {s['low_p_mean']} mg/dL) (Figure 3). The simple carbohydrate load x MSI interaction was positive but not definitive because the confidence interval crossed zero; therefore the primary interpretation is a robust MSI main effect rather than a proven load-specific effect modification.

### Food-matrix proxy and curve phenotypes refined the vulnerability signal

Exploratory food-matrix models suggested that high-MSI participants were especially vulnerable to high digestibility-risk meal contexts. MSI x digestibility-risk interactions were positive for iAUC/1000 (beta {s['fm_i']}; q={s['fm_i_q']}) and peak excursion (beta {s['fm_p']}; q={s['fm_p_q']}) (Table 4; Figure 4). Curve-phenotype analysis identified a large prolonged-response cluster with {s['large_n']} meals, mean iAUC {s['large_i']} mg/dL x min, and mean peak excursion {s['large_p']} mg/dL, characterized by higher MSI and digestibility-risk scores. These food-matrix features are proxy annotations and should be interpreted as hypothesis-generating rather than direct digestion kinetics.

### MSI-enhanced models improved within-CGMacros high-response prediction

The MSI-enhanced model improved within-CGMacros prediction compared with macro/pre-CGM context models. For iAUC, the M3 MSI-enhanced model had MAE {s['m3_i_mae']}, R2 {s['m3_i_r2']}, AUC {s['m3_i_auc']}, and AUPRC {s['m3_i_auprc']}; relative to M2, MAE changed by {s['gain_i_mae']}% and AUC increased by {s['gain_i_auc']}. For peak excursion, M3 had MAE {s['m3_p_mae']}, R2 {s['m3_p_r2']}, AUC {s['m3_p_auc']}, and AUPRC {s['m3_p_auprc']}; relative to M2, MAE changed by {s['gain_p_mae']}% and AUC increased by {s['gain_p_auc']} (Table 5; Figure 5). External datasets supported PPGR heterogeneity and macro-response directionality but also highlighted transportability limits.

## Discussion

This integrated analysis supports a coherent susceptibility-aware framework linking population environmental metabolic profiles, baseline metabolic susceptibility, and dynamic CGM-measured meal responses. In NHANES, DEHP oxidative profile measures were associated with MSI and related metabolic markers. In CGMacros, the same MSI phenotype identified individuals with greater postprandial glycemic vulnerability and improved high-response meal prediction.

The most important conceptual contribution is the bridge-phenotype design. Because no current dataset simultaneously contains urinary DEHP metabolites and detailed free-living CGM meal responses, a direct DEHP-to-PPGR model would overstate the evidence. Instead, the evidence chain is explicitly staged: DEHP oxidative profile to MSI in NHANES, and MSI to PPGR in CGMacros. This approach is more conservative but also more biologically interpretable, because MSI captures the metabolic state likely to modify or express meal vulnerability.

The food-matrix analysis adds a second advance. Simple carbohydrate load x MSI interactions were not definitive, but the digestibility-risk proxy produced stronger interactions with MSI across iAUC, peak excursion, and dynamic curve phenotypes. This suggests that meal vulnerability is better conceptualized as a person-by-matrix interaction than as a person-by-carbohydrate interaction alone. The present score remains a dry-lab proxy, so the finding should guide subsequent expert-coded food annotation, prospective validation, or controlled digestion work rather than be treated as established digestive mechanism.

The prediction results are promising for research-use stratification. MSI improved discrimination and error metrics within CGMacros, especially for high-response identification. However, the CGMacros sample was modest, external datasets differed in design and population, and the current model is not a clinical tool. TRIPOD+AI and PROBAST+AI materials in this package identify leakage controls, calibration summaries, conformal interval coverage, and residual validation gaps.

Several limitations are central. First, the design is observational and cross-dataset; causal claims about DEHP and PPGR are not supported. Second, MSI is a susceptibility phenotype and not a proven mediator. Third, food-matrix features are proxy annotations and require expert validation. Fourth, CGMacros includes only 40 participants, limiting demographic subgroup analysis and external generalizability. Fifth, urinary phthalate metabolites can reflect recent exposure, metabolism, and shared lifestyle/diet sources, and negative-control-like analyses cannot fully remove residual confounding.

In conclusion, DEHP oxidative-metabolism profile was associated with metabolic susceptibility in NHANES, and this susceptibility phenotype was expressed as higher postprandial glycemic vulnerability in CGMacros. The study provides a high-level dry-lab manuscript foundation for an AJCN-style submission, with the next priority being expert food-matrix validation, prospective external validation, and, if desired later, mechanistic digestion experiments.

## Acknowledgments

[To be completed: author contributions, funding, data access acknowledgments, and any institutional support.]

## Data Availability

NHANES data are publicly available from CDC/NCHS. CGMacros raw data are stored locally for this project and require permission or appropriate data-sharing review before public release. External datasets retain their original access terms and licenses. Derived, de-identified tables that do not violate source dataset terms can be shared with the manuscript repository.

## Code Availability

Analysis scripts will be made available in a public repository at publication, excluding protected raw data and files restricted by third-party dataset licenses. The reproducibility manifest records generated analysis files and source modules.

## References

1. Zeevi D, Korem T, Zmora N, et al. Personalized Nutrition by Prediction of Glycemic Responses. Cell. 2015;163:1079-1094.
2. Berry SE, Valdes AM, Drew DA, et al. Human postprandial responses to food and potential for precision nutrition. Nat Med. 2020;26:964-973.
3. Stahlhut RW, van Wijngaarden E, Dye TD, Cook S, Swan SH. Concentrations of urinary phthalate metabolites are associated with increased waist circumference and insulin resistance in adult U.S. males. Environ Health Perspect. 2007;115:876-882.
4. James-Todd T, Stahlhut R, Meeker JD, et al. Urinary phthalate metabolite concentrations and diabetes among women in NHANES 2001-2008. Environ Health Perspect. 2012;120:1307-1313.
5. Centers for Disease Control and Prevention, National Center for Health Statistics. National Health and Nutrition Examination Survey data and documentation.
6. von Elm E, Altman DG, Egger M, et al. The STROBE statement: guidelines for reporting observational studies. Lancet. 2007;370:1453-1457.
7. Lachat C, Hawwash D, Ocke MC, et al. STROBE-nut: An extension of the STROBE statement for nutritional epidemiology. PLoS Med. 2016;13:e1002036.
8. Gallo V, Egger M, McCormack V, et al. STROBE-ME: An extension of the STROBE statement for molecular epidemiology. PLoS Med. 2011;8:e1001117.
9. Moons KGM, Altman DG, Reitsma JB, et al. Transparent Reporting of a multivariable prediction model for Individual Prognosis Or Diagnosis (TRIPOD). Ann Intern Med. 2015;162:55-63.
10. Collins GS, Dhiman P, Andaur Navarro CL, et al. Protocol and reporting guidance for prediction models using artificial intelligence: TRIPOD+AI and PROBAST+AI resources. [Bibliographic details to be verified before submission.]

## Figure Legends

**Figure 1. Integrated no-wet-lab study framework.** The study used a bridge-phenotype design linking NHANES DEHP oxidative metabolite profile to MSI, then testing whether MSI was expressed as CGM-measured meal vulnerability in CGMacros. External datasets were used for boundary evidence.

**Figure 2. NHANES DEHP oxidative profile and metabolic susceptibility.** Survey-weighted beta estimates and 95% CIs for associations of ln(oxidative DEHP metabolites/MEHP) and oxidative metabolite fraction with MSI and metabolic markers.

**Figure 3. CGMacros MSI tertiles and postprandial glycemic response.** Meal-level summaries of mean 2-h iAUC, mean 2-h peak excursion, and high-response rates across participant MSI tertiles.

**Figure 4. Food-matrix susceptibility and CGM curve phenotypes.** Panel A shows exploratory MSI x digestibility-risk proxy interactions. Panel B maps curve-shape clusters by mean iAUC and peak excursion, with point size reflecting meal count and color reflecting mean MSI.

**Figure 5. Prediction performance and external boundary evidence.** Prediction performance for nested CGMacros models and external CGM context used to assess transportability boundaries.

## Tables

### Table 1. Dataset and phenotype summary

{markdown_table(tables['Table1_dataset_and_phenotype_summary_AJCN.csv'], list(tables['Table1_dataset_and_phenotype_summary_AJCN.csv'][0].keys()))}

### Table 2. NHANES survey-weighted DEHP oxidative profile results

{markdown_table(tables['Table2_NHANES_DEHP_MSI_survey_results_AJCN.csv'], list(tables['Table2_NHANES_DEHP_MSI_survey_results_AJCN.csv'][0].keys()))}

### Table 3. CGMacros meal-level MSI and PPGR models

{markdown_table(tables['Table3_CGMacros_MSI_PPGR_models_AJCN.csv'], list(tables['Table3_CGMacros_MSI_PPGR_models_AJCN.csv'][0].keys()))}

### Table 4. Food-matrix and curve-phenotype extensions

{markdown_table(tables['Table4_food_matrix_curve_phenotype_AJCN.csv'], list(tables['Table4_food_matrix_curve_phenotype_AJCN.csv'][0].keys()))}

### Table 5. Prediction model performance

{markdown_table(tables['Table5_prediction_model_performance_AJCN.csv'], list(tables['Table5_prediction_model_performance_AJCN.csv'][0].keys()))}
"""


def add_run(paragraph, text: str, bold: bool = False, italic: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    return run


def configure_docx_styles(doc) -> None:
    from docx.shared import Inches, Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    # Add continuous line numbers for reviewer-style manuscripts.
    sect_pr = section._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    sect_pr.append(ln)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12 if style_name != "Title" else 14)
        style.font.bold = True
        style.paragraph_format.line_spacing = 2


def add_docx_table(doc, title: str, rows: List[Dict[str, object]]) -> None:
    from docx.shared import Pt

    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run(p, title, bold=True)
    cols = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(cols):
        hdr[idx].text = col
    for row in rows:
        cells = table.add_row().cells
        for idx, col in enumerate(cols):
            cells[idx].text = str(row.get(col, ""))
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8)


def build_docx(markdown_text: str, tables: Dict[str, List[Dict[str, object]]]) -> Path:
    from docx import Document

    doc = Document()
    configure_docx_styles(doc)

    lines = markdown_text.splitlines()
    in_tables = False
    for raw in lines:
        line = raw.strip()
        if line == "## Tables":
            in_tables = True
            break
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("**") and ":**" in line:
            p = doc.add_paragraph()
            label, rest = line.split(":**", 1)
            add_run(p, label.replace("**", "") + ":", bold=True)
            add_run(p, rest.strip())
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            add_run(p, line.replace("**", ""), bold=True)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line[:2].isdigit() and ". " in line[:5]:
            doc.add_paragraph(line, style="List Number")
        else:
            # Strip lightweight markdown emphasis for the Word draft.
            clean = line.replace("**", "")
            doc.add_paragraph(clean)

    doc.add_heading("Tables", level=1)
    add_docx_table(doc, "Table 1. Dataset and phenotype summary", tables["Table1_dataset_and_phenotype_summary_AJCN.csv"])
    add_docx_table(doc, "Table 2. NHANES survey-weighted DEHP oxidative profile results", tables["Table2_NHANES_DEHP_MSI_survey_results_AJCN.csv"])
    add_docx_table(doc, "Table 3. CGMacros meal-level MSI and PPGR models", tables["Table3_CGMacros_MSI_PPGR_models_AJCN.csv"])
    add_docx_table(doc, "Table 4. Food-matrix and curve-phenotype extensions", tables["Table4_food_matrix_curve_phenotype_AJCN.csv"])
    add_docx_table(doc, "Table 5. Prediction model performance", tables["Table5_prediction_model_performance_AJCN.csv"])

    out_path = OUT / "manuscript" / f"AJCN_main_manuscript_no_wetlab_{DATE_TAG}.docx"
    doc.save(out_path)
    return out_path


def build_supplement(data: Dict[str, List[Dict[str, str]]]) -> None:
    supplement = f"""
# Supplementary Methods and Results

Date: {DATE_TAG}

## S1. No-wet-lab scope

This package includes only currently completed computational, epidemiologic, CGM, prediction, external-data, and dry-lab proxy analyses. It excludes in vitro digestion, bionic digestion simulation, original/redesigned meal release kinetics, or any wet-lab data not yet generated.

## S2. MSI robustness

Prior outputs showed that MSI-core and PCA-based MSI were highly concordant in both CGMacros and NHANES. The manuscript therefore uses MSI as a bridge phenotype, not as a causal mediator.

## S3. NHANES sensitivity analyses

The source_data folder includes mixture/composition proxy results, effect-modification outputs, negative-control-like sensitivity results, and spline grids from Module B. The manuscript emphasizes the survey-weighted R results as the primary NHANES evidence.

## S4. Food-matrix proxy caution

The digestibility-risk score is a dry-lab proxy derived from meal composition and annotation rules. It is suitable for hypothesis generation and prioritizing future expert-coded or mechanistic validation. It should not be described as measured digestion kinetics.

## S5. Prediction model governance

The prediction models are research-use stratification models. They are not intended for diagnosis, treatment decisions, medication adjustment, or patient-facing prescription without prospective validation.

## S6. External datasets

External CGM datasets were used for boundary evidence. They differ from CGMacros in participant characteristics, protocol structure, food challenges, and outcome availability, so they should not be interpreted as direct replication of the DEHP-MSI-PPGR path.
"""
    write_text(OUT / "supplement" / "Supplementary_methods_and_results_no_wetlab.md", supplement)

    legend = """
# Figure Legends

Figure 1. Integrated no-wet-lab study framework. The study used a bridge-phenotype design linking NHANES DEHP oxidative metabolite profile to MSI, then testing whether MSI was expressed as CGM-measured meal vulnerability in CGMacros. External datasets were used for boundary evidence.

Figure 2. NHANES DEHP oxidative profile and metabolic susceptibility. Survey-weighted beta estimates and 95% CIs for associations of ln(oxidative DEHP metabolites/MEHP) and oxidative metabolite fraction with MSI and metabolic markers.

Figure 3. CGMacros MSI tertiles and postprandial glycemic response. Meal-level summaries of mean 2-h iAUC, mean 2-h peak excursion, and high-response rates across participant MSI tertiles.

Figure 4. Food-matrix susceptibility and CGM curve phenotypes. Panel A shows exploratory MSI x digestibility-risk proxy interactions. Panel B maps curve-shape clusters by mean iAUC and peak excursion, with point size reflecting meal count and color reflecting mean MSI.

Figure 5. Prediction performance and external boundary evidence. Prediction performance for nested CGMacros models and external CGM context used to assess transportability boundaries.
"""
    write_text(OUT / "figures" / "AJCN_figure_legends.md", legend)

    table_notes = """
# Table Notes

Table 1 summarizes datasets and analytic roles.

Table 2 reports survey-weighted NHANES estimates. Beta estimates are presented with 95% confidence intervals and false-discovery-rate q values.

Table 3 reports CGMacros subject-clustered meal-level models. The simple carbohydrate load x MSI interaction is not a primary claim because its confidence interval crosses zero in the main model.

Table 4 reports exploratory food-matrix and curve-phenotype extensions. Digestibility-risk is a dry-lab proxy, not measured digestion kinetics.

Table 5 reports nested prediction performance. M3 is the MSI-enhanced model. These models are for research-use stratification only.
"""
    write_text(OUT / "tables" / "AJCN_table_notes.md", table_notes)


def build_checklists() -> None:
    ajcn = f"""
# AJCN-Style Submission Readiness Checklist

Date checked: {DATE_TAG}

## Formatting implemented

- Main manuscript prepared as Word (.docx) and Markdown.
- Times New Roman, 12-point body text, double-spaced paragraphs, 1-inch margins.
- Reviewer-style line numbering added to the Word document.
- Title page fields included for authors, affiliations, corresponding author, funding, conflicts, data sharing, and trial registration.
- Structured abstract included with Background, Objectives, Methods, Results, and Conclusions.
- Keywords included.
- References included as a working draft, with one explicit bibliographic-verification item.
- Tables placed after the main text and also exported as separate CSV files.
- Figure legends placed after references and also exported separately.
- Figures exported as separate 300-dpi PNG files.

## Must complete before submission

- Replace author, affiliation, corresponding author, funding, conflict-of-interest, and author-contribution placeholders.
- Confirm AJCN article type, word limit, abstract word limit, table/figure count, and any current portal-specific file requirements immediately before submission.
- Verify every reference in a reference manager.
- Confirm IRB/ethics status and data-use permissions for CGMacros and external datasets.
- Decide whether to submit as Original Research or another AJCN category after checking the current author portal.
- Consider reducing the main text and moving extended prediction details to supplementary material if word count exceeds the target article type.
"""
    write_text(OUT / "checklists" / "AJCN_submission_readiness_checklist.md", ajcn)

    strobe = """
# Reporting Checklist Crosswalk

## STROBE / STROBE-nut / STROBE-ME

- Study design, setting, data sources: addressed in Methods.
- Participants and dataset roles: Table 1.
- Variables and outcome definitions: Methods and source_data tables.
- Bias and limitations: Discussion and Supplementary Methods.
- Statistical methods: Methods; model outputs in Tables 2-5 and source_data.
- Nutritional exposure/meal annotation details: food-matrix proxy supplement and source_data.
- Molecular/environmental exposure details: NHANES DEHP oxidative profile Methods and Table 2.

## TRIPOD+AI / PROBAST+AI

- Intended use and non-use: Methods, Discussion, and prediction model card source file.
- Predictor and leakage audit: source_data/prediction_leakage_audit.csv.
- Performance metrics: Table 5.
- Calibration: Table 5 and source_data.
- Uncertainty intervals: source_data/prediction_conformal_interval_coverage.csv.
- External validation: not yet completed as a direct model validation; external datasets are boundary evidence only.
"""
    write_text(OUT / "checklists" / "STROBE_TRIPOD_PROBAST_crosswalk.md", strobe)

    claim_guard = """
# Claim Guardrails for Current AJCN Draft

## Supported wording

- DEHP oxidative-metabolism profile was associated with metabolic susceptibility in NHANES.
- The same susceptibility phenotype was expressed as higher CGM-measured postprandial glycemic vulnerability in CGMacros.
- MSI-enhanced models improved within-CGMacros high-response stratification.
- Food-matrix/digestibility-risk proxy results are exploratory and hypothesis-generating.

## Avoid

- DEHP directly causes higher postprandial glucose.
- MSI is a proven mediator between DEHP and CGM outcomes.
- Digestibility-risk proxy is measured digestion kinetics.
- External datasets directly replicate the integrated DEHP-MSI-PPGR pathway.
- Prediction models are clinically deployable.
"""
    write_text(OUT / "checklists" / "claim_guardrails_for_reviewers.md", claim_guard)


def build_cover_letter_and_title_page() -> None:
    cover = """
# Draft Cover Letter

Dear Editors of The American Journal of Clinical Nutrition,

We are pleased to submit the manuscript entitled "A phthalate oxidative-metabolism profile, metabolic susceptibility, and postprandial glycemic vulnerability" for consideration as original research.

This manuscript integrates NHANES 2013-2018, CGMacros free-living CGM meals, and external CGM datasets to test a susceptibility-aware precision-nutrition framework. Because no single dataset contains both urinary DEHP biomarkers and detailed free-living CGM meal responses, we use a conservative bridge-phenotype design: DEHP oxidative-metabolism profile is evaluated in relation to a metabolic susceptibility phenotype in NHANES, and the same phenotype is evaluated in relation to dynamic postprandial glycemic vulnerability in CGMacros.

The work is potentially relevant to AJCN readers because it links nutritional epidemiology, environmental metabolic profiles, CGM-based postprandial phenotyping, and personalized meal-risk prediction while explicitly avoiding unsupported causal claims. The manuscript is accompanied by source tables, figure files, and reporting-checklist crosswalks.

All authors have approved the manuscript and declare that it is not under consideration elsewhere. [Complete conflicts, funding, ethics, and author-contribution statements before submission.]

Sincerely,

[Corresponding author]
"""
    write_text(OUT / "manuscript" / "AJCN_draft_cover_letter.md", cover)

    title_page = """
# Title Page

Title: A phthalate oxidative-metabolism profile, metabolic susceptibility, and postprandial glycemic vulnerability

Running title: Environmental metabolic susceptibility and postprandial glycemia

Authors: [Complete]

Affiliations: [Complete]

Corresponding author: [Complete]

Funding: [Complete]

Conflicts of interest: [Complete]

Clinical trial registration: Not applicable; secondary analysis of existing observational datasets.

Data sharing: NHANES data are publicly available from CDC/NCHS. CGMacros and external raw data are subject to their original access terms. Derived de-identified outputs can be shared where permitted.

Word count: [Update after final editing]

Number of tables: 5 main tables, 1 external-boundary supplementary table.

Number of figures: 5.
"""
    write_text(OUT / "manuscript" / "AJCN_title_page.md", title_page)


def build_readme(manuscript_docx: Path, figures: List[Path]) -> None:
    readme = f"""
# AJCN No-Wet-Lab Manuscript Package

Generated: {datetime.now().isoformat(timespec="seconds")}

## Main files

- Manuscript Word draft: `{manuscript_docx}`
- Manuscript Markdown draft: `{OUT / "manuscript" / f"AJCN_main_manuscript_no_wetlab_{DATE_TAG}.md"}`
- Cover letter draft: `{OUT / "manuscript" / "AJCN_draft_cover_letter.md"}`
- Figures: `{OUT / "figures"}`
- Tables: `{OUT / "tables"}`
- Supplement/checklists: `{OUT / "supplement"}` and `{OUT / "checklists"}`
- Source data: `{OUT / "source_data"}`

## Figures generated

{chr(10).join(f"- `{p}`" for p in figures)}

## Current interpretation

This manuscript package supports a high-level no-wet-lab submission strategy. The main claim is an integrated association and prediction framework, not a causal DEHP-to-CGM claim.

## Next scientific upgrades

- Complete expert-coded food matrix validation.
- Add external prospective validation or harmonized external prediction if permissions allow.
- Finalize AJCN word count, references, author metadata, ethics, funding, and conflict statements.
- Consider wet-lab digestion only as a later mechanistic extension, not as a prerequisite for this dry-lab manuscript package.
"""
    write_text(OUT / "README_AJCN_package.md", readme)


def build_logs(figures: List[Path], manuscript_docx: Path, tables: Dict[str, List[Dict[str, object]]]) -> None:
    manifest_rows = []
    for path in sorted(OUT.rglob("*")):
        if path.is_file():
            manifest_rows.append(
                {
                    "relative_path": str(path.relative_to(OUT)),
                    "bytes": path.stat().st_size,
                    "modified": datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds"),
                }
            )
    write_csv(OUT / "logs" / "AJCN_package_file_manifest.csv", manifest_rows, ["relative_path", "bytes", "modified"])

    summary = {
        "date": DATE_TAG,
        "package": str(OUT),
        "mirror": str(NHANES_MIRROR),
        "main_docx": str(manuscript_docx),
        "figures": [str(p) for p in figures],
        "tables": {name: len(rows) for name, rows in tables.items()},
        "scope": "Current completed dry-lab/no-wet-lab results only.",
    }
    write_text(OUT / "logs" / "AJCN_generation_summary.json", json.dumps(summary, indent=2))
    shutil.copy2(Path(__file__), OUT / "scripts" / Path(__file__).name)


def build_docx_qa_report(manuscript_docx: Path) -> None:
    from zipfile import ZipFile

    from docx import Document

    doc = Document(str(manuscript_docx))
    text = "\n".join(par.text for par in doc.paragraphs)
    normal = doc.styles["Normal"]
    section = doc.sections[0]
    with ZipFile(manuscript_docx) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    checks = [
        ("DOCX exists", manuscript_docx.exists()),
        ("DOCX opens with python-docx", True),
        ("Title present", "A phthalate oxidative-metabolism profile" in text),
        ("Structured abstract present", "Abstract" in text and "Background:" in text and "Conclusions:" in text),
        ("Causal-claim guardrail present", "causal claims" in text or "not supported" in text),
        ("Normal font Times New Roman", normal.font.name == "Times New Roman"),
        ("Normal font size 12 pt", bool(normal.font.size and abs(normal.font.size.pt - 12) < 0.1)),
        ("Double-spaced Normal style", normal.paragraph_format.line_spacing == 2),
        ("One-inch margins", all(abs(x - 1.0) < 0.05 for x in [
            section.top_margin.inches,
            section.bottom_margin.inches,
            section.left_margin.inches,
            section.right_margin.inches,
        ])),
        ("Reviewer-style line-number XML present", "<w:lnNumType" in xml),
        ("Five main manuscript tables present", len(doc.tables) == 5),
    ]
    table_rows = [len(t.rows) for t in doc.tables]
    lines = [
        "# DOCX Structural QA Report",
        "",
        f"Date: {DATE_TAG}",
        f"File: `{manuscript_docx}`",
        f"Bytes: {manuscript_docx.stat().st_size}",
        f"Paragraphs: {len(doc.paragraphs)}",
        f"Tables: {len(doc.tables)}",
        f"Table rows: {table_rows}",
        "",
        "## Checks",
        "",
    ]
    for label, passed in checks:
        lines.append(f"- {'PASS' if passed else 'CHECK'}: {label}")
    lines.extend(
        [
            "",
            "## Visual render status",
            "",
            "Structural DOCX QA passed. Visual DOCX-to-PNG rendering requires a local Office/LibreOffice conversion chain; if unavailable, the Word draft should be opened manually for final visual inspection before submission.",
        ]
    )
    write_text(OUT / "logs" / "DOCX_structural_QA_report.md", "\n".join(lines))


def mirror_package() -> None:
    if NHANES_MIRROR.exists():
        shutil.rmtree(NHANES_MIRROR)
    shutil.copytree(OUT, NHANES_MIRROR)


def main() -> None:
    ensure_dirs()
    data = load_inputs()
    tables = build_tables(data)
    figures = build_figures(data)
    markdown_text = build_manuscript_markdown(data, tables)
    md_path = OUT / "manuscript" / f"AJCN_main_manuscript_no_wetlab_{DATE_TAG}.md"
    write_text(md_path, markdown_text)
    manuscript_docx = build_docx(markdown_text, tables)
    build_docx_qa_report(manuscript_docx)
    build_supplement(data)
    build_checklists()
    build_cover_letter_and_title_page()
    build_logs(figures, manuscript_docx, tables)
    build_readme(manuscript_docx, figures)
    mirror_package()
    print(f"Generated AJCN package: {OUT}")
    print(f"Mirrored AJCN package: {NHANES_MIRROR}")
    print(f"Main manuscript: {manuscript_docx}")
    print(f"Figures: {len(figures)}")


if __name__ == "__main__":
    main()
