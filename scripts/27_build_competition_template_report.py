from __future__ import annotations

import json
import math
import shutil
from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "competition_template_report_2026-06-11"
FIGPKG = ROOT / "outputs" / "integrated_mainline" / "high_impact_publication_figures_tables_2026-06-11"
NHANES_MIRROR = (
    Path(r"C:\Users\liu12\OneDrive\Desktop\NHANES_MetS_Project")
    / "result"
    / "integrated_mainline"
    / "competition_template_report_2026-06-11"
)
DOCX_OUT = OUT / "Jianxing_Competition_Integrated_DEHP_MSI_PPGR_Paper_2026-06-11.docx"


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def ensure_dirs() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    NHANES_MIRROR.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    table.allow_autofit = False


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_columns(section, num: int, space: int = 360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    if not cols:
        sect_pr.append(cols_el)
    cols_el.set(qn("w:num"), str(num))
    cols_el.set(qn("w:space"), str(space))


def set_a4_template_geometry(section, body: bool = True, columns: int = 1) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    if body:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)
    else:
        section.top_margin = Cm(0.95)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)
    set_columns(section, columns, 360)


def add_page_number(section) -> None:
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3.6)
    normal.paragraph_format.line_spacing = 0.95

    for name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    title = styles["Title"]
    title.font.size = Pt(24)
    title.font.bold = False
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(3)
    title.paragraph_format.space_after = Pt(3)

    h1 = styles["Heading 1"]
    h1.font.size = Pt(10)
    h1.font.bold = False
    h1.font.all_caps = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(4)

    h2 = styles["Heading 2"]
    h2.font.size = Pt(10)
    h2.font.bold = False
    h2.font.italic = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(5)
    h2.paragraph_format.space_after = Pt(2)

    h3 = styles["Heading 3"]
    h3.font.size = Pt(9)
    h3.font.bold = True
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(1)


def p(doc: Document, text: str = "", style: str | None = None, align=None, size: float | None = None, bold: bool | None = None):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    return para


def add_labeled_para(doc: Document, label_text: str, body: str, size: float = 9) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(4)
    r1 = para.add_run(label_text)
    r1.font.name = "Times New Roman"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    r1.font.size = Pt(size)
    r1.font.bold = True
    r2 = para.add_run(body)
    r2.font.name = "Times New Roman"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    r2.font.size = Pt(size)


def add_heading(doc: Document, idx: int, title: str) -> None:
    p(doc, f"{idx}. {title.upper()}", style="Heading 1")


def add_subheading(doc: Document, title: str) -> None:
    p(doc, title, style="Heading 2")


def add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(8)


def add_simple_table(
    doc: Document,
    df: pd.DataFrame,
    caption: str,
    widths: Sequence[int] | None = None,
    font_size: float = 7.4,
    max_rows: int | None = None,
) -> None:
    if max_rows:
        df = df.head(max_rows)
    add_caption(doc, caption)
    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"
    table.allow_autofit = False
    full_width = 9360
    set_table_width(table, full_width)
    if widths is None:
        widths = [int(full_width / cols)] * cols
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        set_cell_width(cell, widths[j])
        set_cell_shading(cell, "E8EEF6")
        set_cell_margins(cell, 80, 80, 80, 80)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(font_size)
                run.font.bold = True
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for j, col in enumerate(df.columns):
            cell = table.cell(i, j)
            val = "" if pd.isna(row[col]) else row[col]
            if isinstance(val, float):
                val = f"{val:.3g}"
            cell.text = str(val)
            set_cell_width(cell, widths[j])
            set_cell_margins(cell, 80, 80, 80, 80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 2 or str(val).replace(".", "", 1).isdigit() else WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                    run.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style=None)
        para.paragraph_format.left_indent = Inches(0.16)
        para.paragraph_format.first_line_indent = Inches(-0.12)
        para.paragraph_format.space_after = Pt(2.2)
        run = para.add_run("\u2022 " + item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        run.font.size = Pt(10)


def add_one_col_section(doc: Document):
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_a4_template_geometry(sec, body=True, columns=1)
    return sec


def add_two_col_section(doc: Document):
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_a4_template_geometry(sec, body=True, columns=2)
    return sec


def load_key_results() -> dict:
    summary_path = OUT / "paper_key_results_summary.json"
    with summary_path.open("r", encoding="utf-8") as f:
        return json.load(f)


def fmt_ci(row: dict, est="Estimate", lo="CI low", hi="CI high") -> str:
    try:
        return f"{float(row[est]):.3f} ({float(row[lo]):.3f}, {float(row[hi]):.3f})"
    except Exception:
        return str(row.get("Estimate (95% CI)", ""))


def build_doc() -> None:
    ensure_dirs()
    key = load_key_results()
    doc = Document()
    configure_styles(doc)
    set_a4_template_geometry(doc.sections[0], body=False, columns=1)
    add_page_number(doc.sections[0])

    title = (
        "An Integrated Dry-Lab Framework Linking DEHP Oxidative-Metabolism Profiles, "
        "Metabolic Susceptibility, and Personalized Postprandial Glycemic Vulnerability"
    )
    p(doc, title, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "ShiYiSiNan CGMacros Project Team", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    p(doc, "Department and Institution to be completed, City, Country", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    p(doc, "Correspondence email to be completed", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    p(doc, "", size=9)

    add_two_col_section(doc)
    add_labeled_para(
        doc,
        "Abstract—",
        (
            "Environmental chemical exposure, metabolic susceptibility, and individualized postprandial glycemic "
            "responses are usually studied in separate datasets. This work reorganized the current project into a "
            "single evidence chain suitable for a high-level undergraduate academic competition: a survey-weighted "
            "NHANES analysis of DEHP oxidative-metabolism profiles, construction and validation of a metabolic "
            "susceptibility index (MSI), free-living CGMacros continuous-glucose-monitoring analyses of meal-level "
            "postprandial glycemic response (PPGR), dry-lab food-matrix annotation, curve phenotype modeling, "
            "prediction calibration, external transportability checks, and a conservative mechanism knowledge map. "
            "The completed computational package contains 9 main figures, 16 extended figures, 5 main tables, "
            "30 supplementary tables, and 69 source-data files. In NHANES, ln(oxidative/MEHP) was associated with "
            "MSI-core (beta=0.183; 95% CI: 0.109, 0.257; q=0.003). In CGMacros, MSI was associated with higher "
            "2-h iAUC (estimate=1.070 per MSI unit for iAUC/1000; 95% CI: 0.540, 1.599; q<0.001). Food-matrix "
            "digestibility-risk and MSI-by-digestibility terms refined carbohydrate-only interpretation, and "
            "MSI-enhanced prediction improved high-response stratification while external datasets defined "
            "generalization boundaries. The study is framed as an observational bridge-phenotype analysis rather "
            "than a direct DEHP-to-PPGR causal mediation study."
        ),
    )
    add_labeled_para(
        doc,
        "Keywords—",
        "NHANES; DEHP; metabolic susceptibility index; CGMacros; continuous glucose monitoring; postprandial glycemic response; dry-lab food matrix; personalized nutrition.",
    )

    add_heading(doc, 1, "Introduction")
    p(
        doc,
        (
            "Precision nutrition research has shown that identical foods can produce highly individualized "
            "postprandial glucose responses. In parallel, population epidemiology suggests that phthalate exposure "
            "and oxidative metabolism may relate to insulin resistance and glycemic traits. The current project "
            "integrates these two evidence layers without overclaiming causality. Its central methodological idea is "
            "to use MSI as a bridge phenotype: DEHP oxidative-metabolism profiles are examined against MSI in NHANES, "
            "and MSI is then tested as a marker of free-living PPGR vulnerability in CGMacros."
        ),
    )
    p(
        doc,
        (
            "The project was deliberately strengthened as a dry-lab-first study. Wet-lab digestive simulation is not "
            "included in the present manuscript version; instead, the computational work emphasizes survey methods, "
            "robust CGM inference, food-matrix annotation, prediction standards, external transportability, and "
            "source-data reproducibility."
        ),
    )

    add_heading(doc, 2, "Study Design and Data Resources")
    p(
        doc,
        (
            "The analysis comprises five connected layers: (1) NHANES 2013-2018 exposure and metabolic phenotype "
            "analysis; (2) MSI construction and sensitivity checks; (3) CGMacros free-living meal-level PPGR modeling; "
            "(4) food-matrix, curve phenotype, counterfactual, and prediction modules; and (5) external boundary and "
            "mechanistic plausibility analyses. The design intentionally separates association, prediction, and "
            "mechanistic plausibility claims."
        ),
    )
    dataset_df = pd.DataFrame(key["dataset"])
    dataset_show = dataset_df.rename(
        columns={
            "Dataset": "Dataset",
            "Analytic role": "Role",
            "Participants": "N",
            "Meals/events": "Events",
            "MSI available": "MSI N",
            "Median MSI": "Median MSI",
        }
    )
    add_one_col_section(doc)
    add_simple_table(
        doc,
        dataset_show[["Dataset", "Role", "N", "Events", "MSI N", "Median MSI"]].head(5),
        "Table 1. Dataset roles and analytic scale.",
        widths=[1500, 3200, 900, 1000, 900, 1100],
        font_size=7.2,
    )
    add_two_col_section(doc)

    add_heading(doc, 3, "Methods")
    add_subheading(doc, "A. Metabolic susceptibility index construction")
    p(
        doc,
        (
            "MSI was built from standardized cardiometabolic components including BMI, HbA1c, fasting glucose, "
            "ln(HOMA-IR), and ln(TG/HDL), using NHANES as the reference scale. Sensitivity versions included complete "
            "component, no-BMI, rank-normalized, and PCA-derived variants. The central requirement was not only "
            "predictive usefulness but robustness across construction variants and datasets."
        ),
    )
    add_subheading(doc, "B. NHANES epidemiology")
    p(
        doc,
        (
            "NHANES analyses evaluated DEHP oxidative-metabolism proxies, especially ln(oxidative/MEHP) and percent "
            "oxidative metabolites, against MSI and glycemic traits using survey-ready results, quartile dose-response "
            "summaries, spline grids, sensitivity checks, effect-modification screens, and mixture/composition proxy "
            "models."
        ),
    )
    add_subheading(doc, "C. CGMacros PPGR, food matrix, and prediction")
    p(
        doc,
        (
            "CGMacros meal-level analyses linked participant MSI with 2-h iAUC, 2-h peak excursion, and high-response "
            "phenotypes under cluster-robust and bootstrap diagnostics. Dry-lab food-matrix features captured available "
            "carbohydrate, GI/GL proxies, refined starch, liquid carbohydrate, visible fiber, protein co-ingestion, "
            "processing markers, and digestibility-risk. Prediction models were evaluated as nested leave-subject-out "
            "models from carbohydrate-only baselines to MSI-enhanced models with calibration and conformal coverage."
        ),
    )

    add_heading(doc, 4, "Results")
    counts = key["counts"]
    p(
        doc,
        (
            f"The final computational display package contains {counts['main_figures']} main figures, "
            f"{counts['extended_figures']} extended figures, {counts['main_tables']} main tables, "
            f"{counts['supp_tables']} supplementary tables, {counts['figure_files']} figure-format files, and "
            f"{counts['source_data']} source-data CSV files."
        ),
    )
    add_subheading(doc, "A. NHANES DEHP oxidative profile and MSI")
    nh_msi = key["nhanes_primary_msi"][0]
    nh_homa = key["nhanes_primary_homa"][0]
    p(
        doc,
        (
            f"In NHANES, ln(oxidative/MEHP) was associated with MSI-core "
            f"(beta={float(nh_msi['Beta']):.3f}; 95% CI: {float(nh_msi['CI low']):.3f}, "
            f"{float(nh_msi['CI high']):.3f}; q={float(nh_msi['q']):.3f}) and ln(HOMA-IR) "
            f"(beta={float(nh_homa['Beta']):.3f}; 95% CI: {float(nh_homa['CI low']):.3f}, "
            f"{float(nh_homa['CI high']):.3f}; q={float(nh_homa['q']):.3f}). Quartile and spline results supported "
            "a monotonic exposure-response pattern for the bridge phenotype."
        ),
    )
    add_subheading(doc, "B. MSI robustness and CGMacros PPGR vulnerability")
    cg_msi = key["cg_msi_iauc"][0]
    p(
        doc,
        (
            f"In CGMacros, MSI was associated with higher 2-h iAUC "
            f"(estimate={float(cg_msi['Estimate']):.3f} for iAUC/1000; 95% CI: "
            f"{float(cg_msi['CI low']):.3f}, {float(cg_msi['CI high']):.3f}; "
            f"q={float(cg_msi['q']):.3g}; {int(cg_msi['Meals'])} meals). Bootstrap and leave-one-subject diagnostics "
            "were used to evaluate whether this signal was driven by a small number of participants."
        ),
    )
    add_subheading(doc, "C. Food-matrix and dynamic curve phenotypes")
    food_int = key["food_interaction"][0]
    p(
        doc,
        (
            f"Food-matrix modeling showed that digestibility-risk and MSI-by-digestibility terms refined a simple "
            f"carbohydrate-only interpretation. For iAUC/1000, the MSI x digestibility-risk term was "
            f"{float(food_int['Estimate']):.3f} (95% CI: {float(food_int['CI low']):.3f}, "
            f"{float(food_int['CI high']):.3f}; q={float(food_int['q']):.3g}). Curve phenotype analysis further "
            "showed that susceptibility can be expressed as delayed, prolonged, or high-peak response shapes, not only "
            "as scalar iAUC or peak values."
        ),
    )
    add_subheading(doc, "D. Prediction, transportability, and mechanism")
    deltas = {d["comparison"]: d for d in key["prediction_delta"]}
    d_i = deltas["M3_msi_enhanced minus M2_macro_precgm_context"]
    p(
        doc,
        (
            f"Adding MSI to the macro/pre-CGM model improved 2-h iAUC prediction by reducing MAE by "
            f"{abs(float(d_i['MAE_delta'])):.1f} mg/dL*min ({abs(float(d_i['MAE_pct_change'])):.1f}%) and increasing "
            f"high-response AUC by {float(d_i['AUC_delta']):.3f}. External CGM datasets were used to define response-scale "
            "and domain-shift boundaries, while a mechanism knowledge graph summarized oxidative stress, insulin "
            "signaling, PPAR signaling, inflammation, mitochondrial function, and beta-cell function pathways."
        ),
    )

    add_one_col_section(doc)
    key_results = pd.DataFrame(
        [
            ["NHANES exposure-response", "ln oxidative/MEHP -> MSI core", f"{float(nh_msi['Beta']):.3f}", f"{float(nh_msi['CI low']):.3f} to {float(nh_msi['CI high']):.3f}", f"{float(nh_msi['q']):.3f}", "Survey-weighted association"],
            ["CGMacros PPGR", "MSI -> 2-h iAUC/1000", f"{float(cg_msi['Estimate']):.3f}", f"{float(cg_msi['CI low']):.3f} to {float(cg_msi['CI high']):.3f}", f"{float(cg_msi['q']):.3g}", "Meal-level vulnerability"],
            ["Food matrix", "MSI x digestibility-risk", f"{float(food_int['Estimate']):.3f}", f"{float(food_int['CI low']):.3f} to {float(food_int['CI high']):.3f}", f"{float(food_int['q']):.3g}", "Person-by-food interaction"],
            ["Prediction", "M3 vs M2 MAE delta", f"{float(d_i['MAE_delta']):.1f}", "", "", "MSI-enhanced stratification"],
            ["Prediction", "M3 vs M2 AUC delta", f"{float(d_i['AUC_delta']):.3f}", "", "", "High-response classification"],
        ],
        columns=["Evidence layer", "Contrast", "Estimate", "95% CI", "q", "Interpretation"],
    )
    add_simple_table(
        doc,
        key_results,
        "Table 2. Condensed primary results across the integrated evidence chain.",
        widths=[1900, 2300, 1050, 1350, 800, 1960],
        font_size=7.4,
    )

    add_heading(doc, 5, "Display Items and Source-Data Package")
    p(
        doc,
        (
            "The current work was consolidated into a reproducible display package. Main figures are inserted below "
            "as full-width displays for readability; extended figures, supplementary tables, and source-data files are "
            "available in the accompanying output directory and indexed in this manuscript."
        ),
    )
    figure_titles = [
        ("Fig. 1.", "Integrated evidence architecture and study design."),
        ("Fig. 2.", "NHANES DEHP oxidative-metabolism profile and metabolic susceptibility."),
        ("Fig. 3.", "Construction, robustness, and transfer of the MSI bridge phenotype."),
        ("Fig. 4.", "MSI identifies free-living postprandial glycemic vulnerability in CGMacros."),
        ("Fig. 5.", "Person-by-food-matrix susceptibility refines carbohydrate-only interpretation."),
        ("Fig. 6.", "CGM curve phenotypes reveal dynamic response vulnerability."),
        ("Fig. 7.", "MSI-enhanced models improve personalized high-response stratification."),
        ("Fig. 8.", "External CGM datasets define transportability boundaries."),
        ("Fig. 9.", "Mechanistic plausibility map linking DEHP oxidative metabolism, MSI, and PPGR vulnerability."),
    ]
    figure_files = [
        FIGPKG / "figures" / "main" / "MainFigure1_EvidenceArchitecture.png",
        FIGPKG / "figures" / "main" / "MainFigure2_NHANES_ExposureResponse.png",
        FIGPKG / "figures" / "main" / "MainFigure3_MSI_Validation.png",
        FIGPKG / "figures" / "main" / "MainFigure4_CGMacros_PPGR_Vulnerability.png",
        FIGPKG / "figures" / "main" / "MainFigure5_FoodMatrix_Interaction.png",
        FIGPKG / "figures" / "main" / "MainFigure6_Dynamic_CGM_Phenotypes.png",
        FIGPKG / "figures" / "main" / "MainFigure7_Prediction_Calibration.png",
        FIGPKG / "figures" / "main" / "MainFigure8_External_Transportability.png",
        FIGPKG / "figures" / "main" / "MainFigure9_Mechanistic_Plausibility.png",
    ]
    for (prefix, cap), img in zip(figure_titles, figure_files):
        if img.exists():
            doc.add_picture(str(img), width=Inches(6.75))
            add_caption(doc, f"{prefix} {cap}")

    add_heading(doc, 6, "Discussion")
    add_two_col_section(doc)
    p(
        doc,
        (
            "This consolidated study is strongest as an integrated computational framework rather than as a single "
            "causal claim. Its novelty lies in joining population environmental epidemiology, bridge-phenotype "
            "construction, free-living CGM physiology, dry-lab food-matrix annotation, prediction standards, and "
            "transportability auditing into one coherent evidence chain. The result is not a claim that DEHP causes "
            "meal-level glycemic excursions in CGMacros, but a cautiously framed model in which DEHP oxidative "
            "metabolism is associated with MSI in NHANES and MSI is expressed as PPGR vulnerability in CGMacros."
        ),
    )
    p(
        doc,
        (
            "The display package also anticipates common reviewer concerns: MSI arbitrariness, survey-vs-Python "
            "reproducibility, small CGMacros sample size, food-matrix proxy validity, prediction leakage, and external "
            "generalizability. These concerns are handled through sensitivity figures, source-data mapping, and "
            "claim-level guardrails."
        ),
    )

    add_heading(doc, 7, "Limitations")
    add_bullets(
        doc,
        [
            "The NHANES component is cross-sectional and should not be interpreted as causal.",
            "MSI is a bridge phenotype, not a proven mediator between DEHP exposure and CGM response.",
            "CGMacros has rich meal-level CGM data but a limited participant count.",
            "Food-matrix and digestibility-risk annotations are dry-lab proxies and require future expert/photo/database validation.",
            "External datasets define transportability boundaries rather than providing exact replication of the full framework.",
            "The current manuscript version intentionally excludes wet-lab digestion experiments.",
        ],
    )

    add_heading(doc, 8, "Conclusion")
    p(
        doc,
        (
            "The project has been reorganized into a coherent high-level competition manuscript: an environmental "
            "exposure signal in NHANES, a validated MSI bridge phenotype, CGMacros PPGR vulnerability, food-matrix "
            "and curve-phenotype refinement, MSI-enhanced prediction, external boundary analysis, and mechanism "
            "plausibility. The completed computational outputs now support a manuscript narrative that is evidence-rich, "
            "transparent, and guarded against causal overstatement."
        ),
    )

    add_one_col_section(doc)
    add_heading(doc, 9, "Supplementary Display Inventory")
    ed_manifest = pd.read_csv(FIGPKG / "logs" / "figure_manifest_high_impact.csv")
    ed = ed_manifest[(ed_manifest["group"] == "extended") & (ed_manifest["format"] == "png")][["figure"]].drop_duplicates()
    ed_titles = [
        "Dataset flow and analytic sample construction",
        "NHANES weighted Table 1 by oxidative quartile",
        "NHANES R survey vs Python reproducibility",
        "NHANES subgroup/effect-modification heatmap",
        "NHANES sensitivity and negative-control-like analyses",
        "DEHP mixture/composition proxy analysis",
        "MSI missingness and version correlations",
        "CGMacros high-response threshold sensitivity",
        "Subject bootstrap and leave-one-subject influence",
        "Food-matrix annotation audit",
        "Counterfactual meal redesign",
        "Prediction leakage audit and predictor dictionary",
        "Prediction conformal interval and net benefit",
        "External domain-shift diagnostics",
        "Mechanism knowledge graph and pathway enrichment",
        "Source-data lineage and reproducibility map",
    ]
    ed_table = pd.DataFrame({"No.": [f"ED Fig. {i}" for i in range(1, 17)], "Display": ed_titles})
    add_simple_table(doc, ed_table, "Table 3. Extended figure inventory.", widths=[1300, 8060], font_size=7.5)

    supp = pd.read_csv(FIGPKG / "tables" / "supplementary" / "supplementary_table_manifest.csv")
    supp_table = supp[["table", "rows", "columns"]].copy()
    supp_table.columns = ["Supplementary table", "Rows", "Columns"]
    add_simple_table(doc, supp_table, "Table 4. Supplementary table inventory.", widths=[6200, 1500, 1500], font_size=6.8, max_rows=30)

    add_heading(doc, 10, "Acknowledgment")
    p(
        doc,
        (
            "Author and affiliation information should be completed before submission. This draft was compiled from "
            "the project's dry-lab computational outputs and follows the provided competition paper template."
        ),
    )

    add_heading(doc, 11, "References")
    refs = key["evidence_refs"]
    for i, ref in enumerate(refs[:12], start=1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.18)
        para.paragraph_format.first_line_indent = Inches(-0.18)
        para.paragraph_format.space_after = Pt(1.6)
        run = para.add_run(f"[{i}] {ref['citation_short']}. {ref['key_message']} {ref['url']}")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        run.font.size = Pt(8)

    doc.save(DOCX_OUT)
    shutil.copy2(DOCX_OUT, NHANES_MIRROR / DOCX_OUT.name)
    print(DOCX_OUT)
    print(NHANES_MIRROR / DOCX_OUT.name)


if __name__ == "__main__":
    build_doc()
